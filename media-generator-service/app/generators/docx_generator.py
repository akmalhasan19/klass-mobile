from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.contracts import DOCX_MIME_TYPE
from app.document_model import RenderDocument, RenderSection, localized_label
from app.generators.base import BaseGenerator, RenderSummary


class DocxGenerator(BaseGenerator):
    export_format = "docx"
    mime_type = DOCX_MIME_TYPE

    def render(self, render_document: RenderDocument, output_path: Path) -> RenderSummary:
        document = Document()
        self._configure_page(document)
        self._configure_typography(document)

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(render_document.title)
        title_run.bold = True
        title_run.font.size = Pt(20)

        summary_paragraph = document.add_paragraph(render_document.summary)
        summary_paragraph.paragraph_format.space_after = Pt(10)

        if render_document.learning_objectives:
            document.add_heading(localized_label(render_document.language, "learning_objectives"), level=1)
            for objective in render_document.learning_objectives:
                document.add_paragraph(objective, style="List Bullet")

        for section in render_document.sections:
            self._add_section(document, section)

        if render_document.activity_blocks:
            document.add_heading(localized_label(render_document.language, "activity_blocks"), level=1)
            for block in render_document.activity_blocks:
                document.add_paragraph(block.title, style="List Bullet")
                detail = document.add_paragraph()
                detail.paragraph_format.left_indent = Inches(0.25)
                detail.add_run(block.instructions)

        document.save(output_path)
        return RenderSummary(page_count=None)

    def _configure_page(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    def _configure_typography(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Aptos"
        normal_style.font.size = Pt(10.5)

    def _add_section(self, document: Document, section: RenderSection) -> None:
        document.add_heading(section.title, level=1)

        for block in section.blocks:
            if block.kind == "paragraph":
                document.add_paragraph(block.content)
                continue

            if block.kind == "bullet":
                document.add_paragraph(block.content, style="List Bullet")
                continue

            prefix = "[ ] " if block.kind == "checklist" else "Note: "
            document.add_paragraph(prefix + block.content)

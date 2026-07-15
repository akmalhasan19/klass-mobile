"""Generate the klass-educational-v1.docx master template (Fase 0B).

This is a build-time helper, not part of the service runtime.  It emits a
``.docx`` master whose body contains ``docxtpl`` Jinja2 placeholders
(``{{ ... }}`` / ``{% ... %}``) so the DOCX engine (Fase 1) can render it from
a ``SlideBlueprint``-derived context without hardcoded formatting.

Run from the media-generator-service root::

    python app/templates/masters/_generate_docx_master.py

The produced file (``klass-educational-v1.docx``) is committed; this script is
kept for reproducibility and can be re-run after designer polish (Fase 6).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

MASTERS_DIR = Path(__file__).resolve().parent
OUTPUT = MASTERS_DIR / "klass-educational-v1.docx"

ACCENT = RGBColor(0x0F, 0x4C, 0x5C)
INK = RGBColor(0x0B, 0x1F, 0x33)


def _set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK


def _add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{ title }}")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = INK

    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subject.add_run("{% if subject %}{{ subject }}"
                           "{% if sub_subject %} · {{ sub_subject }}{% endif %}{% endif %}")
    srun.font.size = Pt(13)
    srun.font.color.rgb = ACCENT
    srun.italic = True

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run("{{ summary }}")


def _add_learning_objectives(doc: Document) -> None:
    doc.add_heading("Tujuan Pembelajaran", level=1)
    doc.add_paragraph("{% for obj in learning_objectives %}• {{ obj }}\n{% endfor %}")


def _add_sections(doc: Document) -> None:
    # docxtpl repeats everything between the lone {% for %} / {% endfor %}
    # paragraphs for each section, so the heading + body below repeat together.
    doc.add_paragraph("{% for section in sections %}")
    doc.add_heading("{{ section.title }}", level=1)
    doc.add_paragraph("{% for block in section.blocks %}"
                      "{% if block.kind == 'bullet' %}• {% endif %}"
                      "{% if block.kind == 'checklist' %}[ ] {% endif %}"
                      "{{ block.content }}\n"
                      "{% endfor %}")
    doc.add_paragraph("{% endfor %}")


def _add_activities(doc: Document) -> None:
    doc.add_heading("Aktivitas dan Penilaian", level=1)
    doc.add_paragraph("{% for activity in activities %}"
                      "{{ activity.title }} — {{ activity.instructions }}\n"
                      "{% endfor %}")


def main() -> None:
    doc = Document()
    _set_base_styles(doc)
    _add_cover(doc)
    _add_learning_objectives(doc)
    _add_sections(doc)
    _add_activities(doc)
    doc.save(str(OUTPUT))
    print(f"Wrote DOCX master -> {OUTPUT}")


if __name__ == "__main__":
    main()

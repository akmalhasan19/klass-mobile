"""Unit tests for ``DocxTemplateEngine`` (Fase 1 / Task 1B).

Tests are organised in three sections:

1. **Context building** (pure logic, no file I/O) — exercise the private
   ``_build_context()`` method which translates a ``SlideBlueprint`` into the
   flat dict consumed by ``docxtpl.DocxTemplate.render()``.

2. **Full rendering** (with the real ``.docx`` master template) — exercise the
   public ``.render()`` method and verify the output is a well-formed,
   ``python-docx``-reopenable file whose paragraph text contains the expected
   content.

3. **Edge cases** — empty sections, no activities, long content, special
   characters, and mixed block kinds.

Design notes
------------
* Context-building tests construct ``SlideBlueprint`` instances programmatically
  (not through ``build_render_document`` / ``build_slide_blueprint``) for precise
  control over every field.  This keeps the test focused on the engine alone and
  avoids coupling to upstream layers.

* Render tests use the real ``klass-educational-v1.docx`` master template from
  ``app/templates/masters/``, mirroring how ``test_template_foundation.py`` uses
  the real HTML master.  Output goes to ``tmp_path`` so no cleanup is needed.

* ``docxtpl`` (being Jinja2-based) auto-escapes content.  Tests therefore
  verify *rendered* text rather than source placeholders.
"""

from __future__ import annotations

from pathlib import Path

import pytest
# pyrefly: ignore [missing-import]
from docx import Document

from app.engines.blueprint import (
    Card,
    ContentBlock,
    DeckMeta,
    Slide,
    SlideBlueprint,
)
from app.engines.docx_template.engine import DocxTemplateEngine

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

_MASTERS_DIR = (
    Path(__file__).resolve().parent.parent
    / "app"
    / "templates"
    / "masters"
)
_MASTER_PATH = _MASTERS_DIR / "klass-educational-v1.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_engine() -> DocxTemplateEngine:
    """Return an engine wired to the real DOCX master template."""
    return DocxTemplateEngine(master_path=_MASTER_PATH)


def _default_deck_meta(**overrides: object) -> DeckMeta:
    """Return a minimal ``DeckMeta`` with overridable defaults."""
    defaults: dict[str, object] = dict(
        title="Handout Pecahan Kelas 5",
        summary="Handout ringkas untuk memperkenalkan pecahan senilai.",
        language="id",
        audience_level="elementary",
        tone="encouraging",
        learning_objectives=[
            "Siswa mengenali pecahan senilai.",
            "Siswa mencoba latihan pecahan dasar.",
        ],
        subject="Matematika",
        sub_subject="Pecahan",
    )
    defaults.update(overrides)  # type: ignore[typeddict-item]
    return DeckMeta(**defaults)  # type: ignore[arg-type]


def _minimal_slides() -> list[Slide]:
    """One content slide + one assessment slide."""
    return [
        Slide(
            slide_type="content",
            title="Tujuan Belajar",
            cards=[
                Card(
                    body_blocks=[
                        ContentBlock(kind="bullet", content="Memahami pecahan senilai"),
                        ContentBlock(kind="bullet", content="Menyelesaikan latihan dasar"),
                    ],
                ),
            ],
        ),
        Slide(
            slide_type="assessment",
            title="Aktivitas dan Penilaian",
            cards=[
                Card(
                    heading="Latihan Mandiri",
                    body_blocks=[
                        ContentBlock(kind="paragraph", content="Kerjakan tiga soal pecahan senilai."),
                    ],
                ),
            ],
        ),
    ]


def _extract_paragraphs(path: Path) -> list[str]:
    """Return a list of non-empty paragraph texts from a ``.docx``."""
    doc = Document(str(path))
    return [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]


# ===========================================================================
# 1. Context building (pure logic, no file I/O)
# ===========================================================================


class TestBuildContext:
    """Exercise ``DocxTemplateEngine._build_context()``.

    These tests create an engine but never call ``.render()`` — they only
    invoke the private context-builder directly, making them pure-logic unit
    tests with zero file-system dependency (beyond importing the engine).
    """

    def test_full_blueprint(self) -> None:
        """All context keys present with expected values."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["title"] == "Handout Pecahan Kelas 5"
        assert ctx["summary"] == "Handout ringkas untuk memperkenalkan pecahan senilai."
        assert ctx["subject"] == "Matematika"
        assert ctx["sub_subject"] == "Pecahan"
        assert ctx["learning_objectives"] == [
            "Siswa mengenali pecahan senilai.",
            "Siswa mencoba latihan pecahan dasar.",
        ]
        assert len(ctx["sections"]) == 1
        assert len(ctx["activities"]) == 1

    def test_sections_contain_title_and_blocks(self) -> None:
        """Each content slide becomes a section dict with title + blocks."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        section = ctx["sections"][0]
        assert section["title"] == "Tujuan Belajar"
        assert len(section["blocks"]) == 2
        assert section["blocks"][0] == {"kind": "bullet", "content": "Memahami pecahan senilai"}
        assert section["blocks"][1] == {"kind": "bullet", "content": "Menyelesaikan latihan dasar"}

    def test_activities_come_from_assessment_slide(self) -> None:
        """Each assessment card becomes an activity entry."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        activity = ctx["activities"][0]
        assert activity["title"] == "Latihan Mandiri"
        assert activity["instructions"] == "Kerjakan tiga soal pecahan senilai."

    def test_activity_fallback_title_when_heading_none(self) -> None:
        """When card.heading is None, the activity title falls back to 'Aktivitas'."""
        slides: list[Slide] = [
            Slide(
                slide_type="assessment",
                title="Aktivitas",
                cards=[
                    Card(
                        heading=None,
                        body_blocks=[
                            ContentBlock(kind="paragraph", content="Do something."),
                        ],
                    ),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["activities"][0]["title"] == "Aktivitas"

    def test_activity_instructions_from_first_block(self) -> None:
        """Instructions are taken from the first body block of each card."""
        slides: list[Slide] = [
            Slide(
                slide_type="assessment",
                title="Aktivitas",
                cards=[
                    Card(
                        heading="Tes",
                        body_blocks=[
                            ContentBlock(kind="paragraph", content="Instruksi utama."),
                            ContentBlock(kind="bullet", content="Poin tambahan."),
                        ],
                    ),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        # Only the *first* block's content becomes the instructions text.
        assert ctx["activities"][0]["instructions"] == "Instruksi utama."

    def test_subject_none(self) -> None:
        """Subject and sub_subject are None when absent from blueprint."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(subject=None, sub_subject=None),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["subject"] is None
        assert ctx["sub_subject"] is None

    def test_empty_learning_objectives(self) -> None:
        """Empty learning_objectives list is forwarded as-is."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(learning_objectives=[]),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["learning_objectives"] == []

    def test_only_title_slide_no_sections_no_activities(self) -> None:
        """Blueprint with only a title slide yields empty sections and activities."""
        slides: list[Slide] = [
            Slide(
                slide_type="title",
                title="Judul",
                cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content="Halo")])],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["sections"] == []
        assert ctx["activities"] == []

    def test_only_content_slides_no_assessment(self) -> None:
        """Blueprint with content slides but no assessment yields empty activities."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Bagian 1",
                cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content="Teks.")])],
            ),
            Slide(
                slide_type="content",
                title="Bagian 2",
                cards=[Card(body_blocks=[ContentBlock(kind="bullet", content="Item.")])],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert len(ctx["sections"]) == 2
        assert ctx["activities"] == []

    def test_multiple_content_slides_produce_multiple_sections(self) -> None:
        """Each content slide maps to exactly one section entry."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title=f"Bagian {i}",
                cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content=f"Konten {i}")])],
            )
            for i in range(5)
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert len(ctx["sections"]) == 5
        for i, section in enumerate(ctx["sections"]):
            assert section["title"] == f"Bagian {i}"

    def test_content_slide_with_multiple_cards_all_blocks_aggregated(self) -> None:
        """All cards' body blocks are flattened into a single blocks list."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Multi-Card",
                cards=[
                    Card(body_blocks=[ContentBlock(kind="paragraph", content="Card1")]),
                    Card(body_blocks=[ContentBlock(kind="bullet", content="Card2 B1"),
                                      ContentBlock(kind="bullet", content="Card2 B2")]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert len(ctx["sections"][0]["blocks"]) == 3

    def test_mixed_block_kinds(self) -> None:
        """Different block kinds (paragraph, bullet, checklist, note) are preserved."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Mixed",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content="Paragraf satu."),
                        ContentBlock(kind="bullet", content="Poin pertama."),
                        ContentBlock(kind="checklist", content="Ceklis ini."),
                        ContentBlock(kind="note", content="Catatan penting."),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        blocks = ctx["sections"][0]["blocks"]
        assert len(blocks) == 4
        assert blocks[0] == {"kind": "paragraph", "content": "Paragraf satu."}
        assert blocks[1] == {"kind": "bullet", "content": "Poin pertama."}
        assert blocks[2] == {"kind": "checklist", "content": "Ceklis ini."}
        assert blocks[3] == {"kind": "note", "content": "Catatan penting."}


# ===========================================================================
# 2. Full rendering (real master template + file I/O)
# ===========================================================================


class TestRender:
    """Exercise the full ``render()`` pipeline with the real master template.

    These tests create actual ``.docx`` files on disk and then verify them
    with ``python-docx``.
    """

    def test_renders_reopenable_docx(self, tmp_path: Path) -> None:
        """Output file is a valid ``.docx`` that ``python-docx`` can open."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)

        assert output.exists()
        assert output.stat().st_size > 0
        # Must be reopenable without exception.
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_render_contains_title(self, tmp_path: Path) -> None:
        """The deck title appears in the rendered paragraph text."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = _extract_paragraphs(output)

        assert any("Handout Pecahan Kelas 5" in t for t in texts)

    def test_render_contains_summary(self, tmp_path: Path) -> None:
        """Summary text is rendered in the document."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = _extract_paragraphs(output)

        assert any("Handout ringkas" in t for t in texts)

    def test_render_contains_learning_objectives(self, tmp_path: Path) -> None:
        """Learning objective bullets appear in the output."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = "\n".join(_extract_paragraphs(output))

        assert "Siswa mengenali pecahan senilai" in texts
        assert "Siswa mencoba latihan pecahan dasar" in texts

    def test_render_contains_section_content(self, tmp_path: Path) -> None:
        """Content section bullet text appears in the rendered output."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = "\n".join(_extract_paragraphs(output))

        assert "Memahami pecahan senilai" in texts

    def test_render_contains_activity_content(self, tmp_path: Path) -> None:
        """Activity block text appears in the rendered output."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = "\n".join(_extract_paragraphs(output))

        assert "Latihan Mandiri" in texts
        assert "Kerjakan tiga soal" in texts

    def test_render_subject_and_sub_subject(self, tmp_path: Path) -> None:
        """Subject and sub-subject labels appear in the rendered output."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "output.docx"

        _make_engine().render(blueprint, output)
        texts = "\n".join(_extract_paragraphs(output))

        # The master template renders them as "{{ subject }} · {{ sub_subject }}".
        # Both "Matematika" and "Pecahan" should appear.
        assert "Matematika" in texts
        assert "Pecahan" in texts


# ===========================================================================
# 3. Edge cases
# ===========================================================================


class TestEdgeCases:
    """Context building for unusual or extreme inputs."""

    def test_long_content(self) -> None:
        """A very long content string is not truncated in the context dict."""
        long_text = "X" * 1800  # within the 2000-byte max_length
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Panjang",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content=long_text),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["sections"][0]["blocks"][0]["content"] == long_text

    def test_long_content_renders_without_error(self, tmp_path: Path) -> None:
        """A blueprint with long content renders without crashing."""
        long_text = "Y" * 1800
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Panjang",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content=long_text),
                        ContentBlock(kind="bullet", content="Poin normal"),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        output = tmp_path / "long.docx"

        _make_engine().render(blueprint, output)

        assert output.exists()
        texts = "\n".join(_extract_paragraphs(output))
        assert "Poin normal" in texts

    def test_unicode_content(self) -> None:
        """Unicode characters are preserved in the context dict."""
        unicode_text = "Pecahan ⅓ dan ¼ — pelajaran ke-5 (matematika dasar)"
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Unicode",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content=unicode_text),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["sections"][0]["blocks"][0]["content"] == unicode_text

    def test_unicode_content_renders(self, tmp_path: Path) -> None:
        """Unicode characters survive the render-and-reopen round-trip."""
        unicode_text = "Pecahan ⅓ dan ¼ — pelajaran ke-5"
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Unicode",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content=unicode_text),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        output = tmp_path / "unicode.docx"

        _make_engine().render(blueprint, output)

        texts = "\n".join(_extract_paragraphs(output))
        assert "⅓" in texts
        assert "¼" in texts

    def test_special_chars_in_title(self) -> None:
        """Special characters in the deck title survive the context builder."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(
                title='Matematika "Pecahan" & <Lanjutan>',
            ),
            slides=_minimal_slides(),
        )
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["title"] == 'Matematika "Pecahan" & <Lanjutan>'

    def test_special_chars_renders(self, tmp_path: Path) -> None:
        """Special characters in content survive render and reopen."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Spesial",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content='Harga Rp 5.000 < 10% diskon > "MURAH!"'),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=slides,
        )
        output = tmp_path / "special.docx"

        _make_engine().render(blueprint, output)

        texts = "\n".join(_extract_paragraphs(output))
        assert "Rp" in texts
        assert "diskon" in texts

    def test_newlines_in_content(self) -> None:
        """Newline characters in content are preserved."""
        multiline = "Baris pertama\nBaris kedua\nBaris ketiga"
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title="Multiline",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind="paragraph", content=multiline),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert ctx["sections"][0]["blocks"][0]["content"] == multiline

    def test_no_title_slide_skipped(self) -> None:
        """Title slides are skipped in section mapping (not content or assessment)."""
        slides: list[Slide] = [
            Slide(
                slide_type="title",
                title="Judul",
                cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content="Skip")])],
            ),
            Slide(
                slide_type="content",
                title="Mulai",
                cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content="Konten")])],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        ctx = _make_engine()._build_context(blueprint)

        assert len(ctx["sections"]) == 1
        assert ctx["sections"][0]["title"] == "Mulai"

    def test_render_multiple_activities(self, tmp_path: Path) -> None:
        """Multiple assessment cards produce multiple activities in output."""
        slides: list[Slide] = [
            Slide(
                slide_type="assessment",
                title="Aktivitas",
                cards=[
                    Card(
                        heading="Aktivitas 1",
                        body_blocks=[ContentBlock(kind="paragraph", content="Kerjakan A.")],
                    ),
                    Card(
                        heading="Aktivitas 2",
                        body_blocks=[ContentBlock(kind="paragraph", content="Kerjakan B.")],
                    ),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        output = tmp_path / "multi-activity.docx"

        _make_engine().render(blueprint, output)

        texts = "\n".join(_extract_paragraphs(output))
        assert "Aktivitas 1" in texts
        assert "Aktivitas 2" in texts

    def test_render_no_assessment(self, tmp_path: Path) -> None:
        """Blueprint with only content slides renders activities section as empty."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=[
                Slide(
                    slide_type="content",
                    title="Hanya Konten",
                    cards=[
                        Card(body_blocks=[
                            ContentBlock(kind="paragraph", content="Teks konten."),
                        ]),
                    ],
                ),
            ],
        )
        output = tmp_path / "no-assessment.docx"

        _make_engine().render(blueprint, output)

        assert output.exists()
        texts = "\n".join(_extract_paragraphs(output))
        assert "Teks konten" in texts

    def test_render_no_content_slides_at_all(self, tmp_path: Path) -> None:
        """Blueprint with only a title slide does not crash the renderer."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(title="Only Title"),
            slides=[
                Slide(
                    slide_type="title",
                    title="Only Title",
                    cards=[Card(body_blocks=[ContentBlock(kind="paragraph", content="Just title.")])],
                ),
            ],
        )
        output = tmp_path / "only-title.docx"

        # Should not raise — the template gracefully handles empty loops.
        _make_engine().render(blueprint, output)

        assert output.exists()
        texts = "\n".join(_extract_paragraphs(output))
        assert "Only Title" in texts

    def test_magic_bytes(self, tmp_path: Path) -> None:
        """Rendered file starts with the ZIP magic bytes (DOCX = ZIP)."""
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        output = tmp_path / "magic.docx"

        _make_engine().render(blueprint, output)

        raw = output.read_bytes()
        assert raw[:2] == b"PK"  # ZIP magic bytes

    @pytest.mark.parametrize("block_kind", ["paragraph", "bullet", "checklist", "note"])
    def test_each_block_kind_renders_without_error(self, block_kind: str, tmp_path: Path) -> None:
        """Every supported ``ContentBlock.kind`` renders without a crash."""
        slides: list[Slide] = [
            Slide(
                slide_type="content",
                title=f"Kind {block_kind}",
                cards=[
                    Card(body_blocks=[
                        ContentBlock(kind=block_kind, content=f"Test konten {block_kind}."),
                    ]),
                ],
            ),
        ]
        blueprint = SlideBlueprint(deck_meta=_default_deck_meta(), slides=slides)
        output = tmp_path / f"{block_kind}.docx"

        _make_engine().render(blueprint, output)

        assert output.exists()
        texts = "\n".join(_extract_paragraphs(output))
        assert f"Test konten {block_kind}" in texts

    def test_render_idempotent(self, tmp_path: Path) -> None:
        """Rendering the same blueprint twice produces identical paragraph text.

        We compare extracted paragraph text rather than raw bytes because
        ``python-docx`` writes a ``modified`` timestamp into ``core.xml`` on
        every ``save()``, making byte-wise equality unreliable.
        """
        blueprint = SlideBlueprint(
            deck_meta=_default_deck_meta(),
            slides=_minimal_slides(),
        )
        engine = _make_engine()
        out1 = tmp_path / "run1.docx"
        out2 = tmp_path / "run2.docx"

        engine.render(blueprint, out1)
        engine.render(blueprint, out2)

        assert _extract_paragraphs(out1) == _extract_paragraphs(out2)
